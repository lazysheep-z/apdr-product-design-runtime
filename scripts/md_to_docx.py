#!/usr/bin/env python3
"""Convert 工作交接文档.md to Word (.docx)."""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill: str = "E7E6E6"):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def strip_md_inline(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"_([^_]+)_", r"\1", text)
    return text


def add_rich_paragraph(doc, text: str, style=None, italic=False):
    p = doc.add_paragraph(style=style)
    if not text.strip():
        return p
    parts = re.split(r"(\*\*[^*]+\*\*|\[[^\]]+\]\([^)]+\))", text)
    for part in parts:
        if not part:
            continue
        m_link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
        m_bold = re.match(r"\*\*([^*]+)\*\*", part)
        if m_link:
            run = p.add_run(m_link.group(1))
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.underline = True
        elif m_bold:
            run = p.add_run(m_bold.group(1))
            run.bold = True
        else:
            run = p.add_run(part)
            if italic:
                run.italic = True
    return p


def parse_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [c.strip() for c in line.split("|")]


def is_separator_row(cells: list[str]) -> bool:
    return all(re.match(r"^:?-+:?$", c.replace(" ", "")) or c == "" for c in cells)


def add_table(doc, rows: list[list[str]]):
    if len(rows) < 2:
        return
    header = [strip_md_inline(c) for c in rows[0]]
    body_start = 2 if len(rows) > 1 and is_separator_row(rows[1]) else 1
    body = [[strip_md_inline(c) for c in r] for r in rows[body_start:]]
    cols = len(header)
    table = doc.add_table(rows=1 + len(body), cols=cols)
    table.style = "Table Grid"
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_shading(cell)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for i, row in enumerate(body):
        for j in range(cols):
            val = row[j] if j < len(row) else ""
            table.rows[i + 1].cells[j].text = val
    doc.add_paragraph()


def convert(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "PingFang SC"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    style.font.size = Pt(11)

    i = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    table_buf: list[list[str]] = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            add_table(doc, table_buf)
            table_buf = []

    while i < len(lines):
        line = lines[i]

        if in_code:
            if line.strip().startswith("```"):
                text = "\n".join(code_lines)
                if code_lang == "mermaid":
                    add_rich_paragraph(doc, "【流程图示意，建议在 Word 中插入 SmartArt 或 Visio】", italic=True)
                    p = doc.add_paragraph(style="No Spacing")
                    run = p.add_run(text)
                    run.font.size = Pt(9)
                    run.italic = True
                else:
                    p = doc.add_paragraph(style="No Spacing")
                    run = p.add_run(text)
                    run.font.name = "Menlo"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
                    run.font.size = Pt(9)
                doc.add_paragraph()
                in_code = False
                code_lines = []
                code_lang = ""
            else:
                code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith("```"):
            flush_table()
            in_code = True
            code_lang = line.strip()[3:].strip()
            i += 1
            continue

        if line.strip().startswith("|"):
            table_buf.append(parse_table_row(line))
            i += 1
            continue

        flush_table()

        if line.strip() == "---":
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(strip_md_inline(line[2:]), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(strip_md_inline(line[3:]), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(strip_md_inline(line[4:]), level=2)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(strip_md_inline(line[5:]), level=3)
            i += 1
            continue

        if line.startswith("> "):
            add_rich_paragraph(doc, line[2:], style="Intense Quote")
            i += 1
            continue

        if re.match(r"^(\s*)- ", line):
            indent = len(line) - len(line.lstrip())
            text = line.strip()[2:]
            style = "List Bullet" if indent < 2 else "List Bullet 2"
            add_rich_paragraph(doc, text, style=style)
            i += 1
            continue

        if line.strip().startswith("**") and line.strip().endswith("**") and "|" not in line:
            add_rich_paragraph(doc, line.strip())
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        add_rich_paragraph(doc, line)
        i += 1

    flush_table()
    doc.save(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md = root / "docs" / "工作交接文档.md"
    out = root / "docs" / "工作交接文档.docx"
    if len(sys.argv) > 1:
        md = Path(sys.argv[1])
        out = Path(sys.argv[2]) if len(sys.argv) > 2 else md.with_suffix(".docx")
    convert(md, out)
